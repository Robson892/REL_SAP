import pandas as pd
import streamlit as st
from time import sleep
from datetime import datetime, timedelta
import locale
from PIL import Image
import io   
from lxml import html
import sys

def sap():
    import streamlit as st
    # Configurar o título da aba e o ícone
    st.set_page_config(page_title="Gerador de Relatório SAP", page_icon=":rocket:")


    # Carregando o logo
    logo = Image.open('LOGO AG.png')

    # Exibindo o logo
    st.image(logo, width=150)


    with open("style.css") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
        

    #Titulo so site
    st.title("Gerador de Relatório SAP")

    def obter_periodo():
        periodo_inicio_dt = None
        periodo_final_dt = None

        for i in range(2):  # Apenas duas iterações para os dois períodos
            try:
                # Obter o período inserido pelo usuário
                periodo = st.text_input(f"Insira o {'primeiro' if i == 0 else 'segundo'} período (no formato DD/MM/YYYY): ", key=f"periodo_{i}")

                # Remover barras extras que o usuário pode tentar inserir manualmente
                periodo = periodo.replace("/", "")

                # Verificar se a entrada do período está vazia
                if not periodo:
                    st.warning(f"Por favor, insira o {'primeiro' if i == 0 else 'segundo'} período.")
                    return None, None

                # Converter a entrada do período para um objeto datetime
                periodo_dt = datetime.strptime(periodo, "%d%m%Y")

                if i == 0:
                    periodo_inicio_dt = periodo_dt
                else:
                    periodo_final_dt = periodo_dt

            except ValueError:
                st.error("Formato de data inválido. Por favor, insira as datas no formato DD/MM/YYYY.")
                return None, None

        # Verificar se o período final é posterior ao período inicial
        if periodo_final_dt and periodo_inicio_dt and periodo_final_dt < periodo_inicio_dt:
            st.error("O período final deve ser posterior ao período inicial.")
            return None, None

        # Formatar os períodos como strings no formato DD/MM/YYYY
        periodo_inicio_str = periodo_inicio_dt.strftime('%d/%m/%Y') if periodo_inicio_dt else None
        periodo_final_str = periodo_final_dt.strftime('%d/%m/%Y') if periodo_final_dt else None

        return periodo_inicio_str, periodo_final_str

    # Obter o período do usuário com tratamento de erros
    periodo_inicio, periodo_final = obter_periodo()

    if periodo_inicio and periodo_final:
        # Exibir os períodos formatados
        st.success(f"Período inicial informado: {periodo_inicio}")
        st.success(f"Período final informado: {periodo_final}")

    # Definir o idioma como português, com tratamento de erro
    try:
        if sys.platform == 'win32':
            locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')
        else:
            locale.setlocale(locale.LC_TIME, 'pt_BR.utf8')
        #st.write("Locale definido com sucesso.")
    except locale.Error:
        st.error("Não foi possível definir o locale para 'Portuguese_Brazil.1252'.")

    # Obter a data de hoje
    hoje = datetime.today()
    st.write(f"Data de hoje: {hoje.strftime('%d/%m/%Y')}")

    # Mapeamento dos nomes dos dias da semana em português
    dias_da_semana = {
        0: 'Segunda',
        1: 'Terça',
        2: 'Quarta',
        3: 'Quinta',
        4: 'Sexta',
        5: 'Sabado',
        6: 'Domingo'
    }

    # Subtrair um dia para obter a data de ontem
    ontem = hoje - timedelta(days=1)
    # Obter o dia da semana de ontem (0 = segunda-feira, 1 = terça-feira, ..., 6 = domingo)
    dia_da_semana_de_ontem = ontem.weekday()

    # Obter o nome do dia da semana de ontem usando o mapeamento
    dia_da_semana_de_ontem_formatado = dias_da_semana[dia_da_semana_de_ontem]

    st.success(f'Dia da semana: {dia_da_semana_de_ontem_formatado}')
    # Formatar o dia de ontem
    dia_mes = "Dia " + ontem.strftime('%d')
    st.success(f'Dia do mês: {dia_mes}')


    banco_azp = st.file_uploader("Selecione um arquivo Excel (AZP):", type="xlsx", key="banco_azp")
    banco_bip = st.file_uploader("Selecione um arquivo Excel (BIP):", type="xlsx", key="banco_bip")
    banco_fip = st.file_uploader("Selecione um arquivo Excel (FIP):", type="xlsx", key="banco_fip")
    banco_grp = st.file_uploader("Selecione um arquivo Excel (GRP):", type="xlsx", key="banco_grp")
    banco_pip = st.file_uploader("Selecione um arquivo Excel (PIP):", type="xlsx", key="banco_pip")

    # CSS para estilizar e alterar o texto dos componentes
    st.markdown("""
    <style>
        <style>
        /* Estilizar o botão de upload */
        .st-eb button {
            background-color: #4CAF50; /* Cor de fundo verde */
            color: white; /* Texto branco */
            font-size: 16px; /* Tamanho da fonte */
        }
        
        /* Alterar o texto do botão de upload */
        .st-eb button::before {
            content: 'Selecione um arquivo';
            white-space: pre;
        }

        /* Alterar o texto da área de dropzone */
        div[data-testid="stFileUploaderDropzoneInstructions"] > div {
            font-size: 16px; /* Tamanho da fonte */
            color: black; /* Cor do texto */
        }
        div[data-testid="stFileUploaderDropzoneInstructions"] > div::before {
            content: 'Arraste e solte o arquivo aqui';
            white-space: pre;
            font-size: 16px; /* Tamanho da fonte */
            color: black; /* Cor do texto */
        }
        div[data-testid="stFileUploaderDropzoneInstructions"] > div > span {
            display: none;
        }
            
    
        </style>
    """, unsafe_allow_html=True)

    # Exibir os arquivos carregados (opcional)
    if banco_azp is not None:
        st.write("Arquivo AZP carregado com sucesso!")
    if banco_bip is not None:
        st.write("Arquivo BIP carregado com sucesso!")
    if banco_fip is not None:
        st.write("Arquivo FIP carregado com sucesso!")
    if banco_grp is not None:
        st.write("Arquivo GRP carregado com sucesso!")
    if banco_pip is not None:
        st.write("Arquivo PIP carregado com sucesso!")

    if st.button("Iniciar"):
        

            
            
        st.write("Processo iniciado, gerando documento!")   

        if dia_da_semana_de_ontem_formatado == dia_da_semana_de_ontem_formatado:

            

            if dia_mes == dia_mes:
                    
                
                from docx import Document
                from docx.shared import Pt
                from docx.enum.table import WD_ALIGN_VERTICAL
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                from docx.shared import RGBColor


                def aplicar_formatacao(run):
                    run.font.name = 'Calibri Light'
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Cor preta
                    run.bold = False
                    run.italic = False
                    #run.underline = False

                # Crie um novo documento do Word
                doc = Document()

                # Acesse a seção do documento (por padrão, há pelo menos uma seção)
                section = doc.sections[0]

                # Acesse o cabeçalho da seção (por padrão, há um cabeçalho em branco)
                header = section.header

                # Adicione texto ao cabeçalho
                header.paragraphs[0].text = 'Monitoramento de Dumps e Jobs dos Ambientes SAP'
                header.paragraphs[0].runs[0].font.size = Pt(14)  # Tamanho da fonte
                header.paragraphs[0].runs[0].bold = True  # Texto em negrito
                cabeçalho_paragrafo = header.paragraphs[0]
                cabeçalho_paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Configure as margens estreitas (0,5 polegadas) para todas as páginas
                sections = doc.sections
                for section in sections:
                    section.left_margin = Pt(36)  # 0,5 polegadas
                    section.right_margin = Pt(36)  # 0,5 polegadas
                    section.top_margin = Pt(36)  # 0,5 polegadas
                    section.bottom_margin = Pt(36)  # 0,5 polegadas

                ## Função para adicionar uma tabela a uma seção específica do documento
                def add_dataframe_as_table(dataframe, section_title):
                    doc.add_heading(section_title, level=1)

                    table = doc.add_table(rows=1, cols=len(dataframe.columns))
                    table.style = 'Table Grid'

                    for col_num, column_name in enumerate(dataframe.columns):
                        table.cell(0, col_num).text = column_name
                        cell = table.cell(0, col_num)
                        cell.paragraphs[0].alignment = 1  # Centraliza o cabeçalho
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        # Defina o tamanho da fonte para tamanho 5 (5 pt) nas células do DataFrame
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(5)

                    for row in dataframe.itertuples(index=False):
                        cells = table.add_row().cells
                        for col_num, cell_value in enumerate(row):
                            
                            cells[col_num].text = str(cell_value)
                            # Defina o tamanho da fonte para tamanho 5 (5 pt) nas células do DataFrame
                            for paragraph in cells[col_num].paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(5)

                    # Configure a tabela para ajustar automaticamente ao conteúdo
                    table.allow_autofit = True
                # Adicione os DataFrames como tabelas ao documento
                sumario = [
                f'1. AZP',
                f'1.1. Jobs cancelados',
                f'1.2. Jobs previstos e não executados',
                f'1.3. Desvio de ocorrências previstas x executadas',
                f'1.4. Jobs ativos duplicados',
                f'1.5. Jobs ativos acima de 24h',
                f'1.6. Jobs com tempo de execução diferente do previsto',
                f'1.7. Jobs com atraso no início da execução (>10min)',
                f'1.8. Dumps - AZP',
                f'2. FIP',
                f'2.1. Jobs cancelados',
                f'2.2. Jobs previstos e não executados',
                f'2.3. Desvio de ocorrências previstas x executadas',
                f'2.4. Jobs ativos duplicados',
                f'2.5. Jobs ativos acima de 24h',
                f'2.6. Jobs com tempo de execução diferente do previsto',
                f'2.7. Jobs com atraso no início da execução (>10min)',
                f'2.8. Dumps - FIP',
                f'3. PIP',
                f'3.1. Jobs cancelados',
                f'3.2. Jobs previstos e não executados',
                f'3.3. Desvio de ocorrências previstas x executadas',
                f'3.4. Jobs ativos duplicados',
                f'3.5. Jobs ativos acima de 24h',
                f'3.6. Jobs com tempo de execução diferente do previsto',
                f'3.7. Jobs com atraso no início da execução (>10min)',
                f'3.8. Dumps - BIP',
                f'4. GRP',
                f'4.1. Jobs cancelados',
                f'4.2. Jobs previstos e não executados',
                f'4.3. Desvio de ocorrências previstas x executadas',
                f'4.4. Jobs ativos duplicados',
                f'4.5. Jobs ativos acima de 24h',
                f'4.6. Jobs com tempo de execução diferente do previsto',
                f'4.7. Jobs com atraso no início da execução (>10min)',
                f'4.8. Dumps - GRP',
                f'5. BIP',
                f'5.1. Jobs cancelados',
                f'5.2. Jobs previstos e não executados',
                f'5.3. Desvio de ocorrências previstas x executadas',
                f'5.4. Jobs ativos duplicados',
                f'5.5. Jobs ativos acima de 24h',
                f'5.6. Jobs com tempo de execução diferente do previsto',
                f'5.7. Jobs com atraso no início da execução (>10min)'
                f'5.8. Dumps - BIP',

                ]

                #----------AMBIENTE AZP---------#
                
                # Tabelas e filtros por base
                bd_azp1 = pd.read_excel(banco_azp ,sheet_name = 'Jobs')
                bd_azp2 = pd.read_excel(banco_azp ,sheet_name = 'Repetitivo')
                bd_azp3 = pd.read_excel(banco_azp ,sheet_name = 'Diário')
                bd_azp4 = pd.read_excel(banco_azp ,sheet_name = 'Semanal')
                bd_azp5 = pd.read_excel(banco_azp ,sheet_name = 'Mensal')
                bd_azp6 = pd.read_excel(banco_azp ,sheet_name = 'Ativo_24h')
                bd_azp7 = pd.read_excel(banco_azp, sheet_name = 'DUMPS' )

                # filtros por colunas
                azp2 = bd_azp2.iloc[ : , 0:11]
                azp3 = bd_azp3.iloc[ : , 0:9]
                azp4 = bd_azp4.iloc[ : , 0:10]
                azp5 = bd_azp5.iloc[ : , 0:10]
                azp6 = bd_azp6.iloc[0:8]        
                    
                #1.1 Jobs Cancelados
                cancelado = bd_azp1.loc[bd_azp1['Status'] == 'cancelado']
                df = pd.DataFrame(cancelado)
                df_cancelado = df.dropna(how='all', axis=0, inplace=False)


                # 1.2 Jobs Previstos e não executados
                #Repetitvo
                excluir = [3,4,5]       
                repetitivo = azp2.iloc[:, :11][azp2['Ocorrencias Reais'] == 0]
                df = pd.DataFrame(repetitivo)
                df = df.drop(df.columns[excluir], axis=1)
                df_repetitivo = df.dropna(how='all', axis=0, inplace=False)

                #Diario
                diario = azp3.iloc[:, :10][azp3['Status'] == 'cancelado']
                df = pd.DataFrame(diario)
                db_diario = df.dropna(how='all', axis=0, inplace=False)

                #semanal
                semana = azp4.iloc[:, :10][azp4['Dia da Semana'] == dia_da_semana_de_ontem_formatado]
                df = pd.DataFrame(semana)
                df_semana = df.dropna(how='all', axis=0, inplace=False)

                #Mensal
                mensal = azp5.iloc[:, :9][azp5['Dia do Mes'] == dia_mes]
                df = pd.DataFrame(mensal)
                df_mensal = df.dropna(how='all', axis=0, inplace=False)

                #1.3 Desvio de ocorrências previstas x executadas
                
            
                repetitivo_0 = azp2.iloc[:, :11][azp2['Desvio Tempo em %'] != 0]
                df = pd.DataFrame(repetitivo_0)

                df = df.drop(df.columns[excluir], axis=1)#excluir
                
                desvio_ocorrencias = df.dropna(how='all', axis=0, inplace=False)

                #1.4 Jobs Ativos Duplicados
                df = pd.DataFrame(azp6)
                #df_ativo = df.dropna(how='all', axis=0, inplace=False)

                #1.6.	Jobs com tempo de execução diferente do previsto

                semana_1 = azp4[(azp4['Desvio Tempo em %'] < -30) | (azp4['Desvio Tempo em %'] > 30)]
                mensal_1 = azp5[(azp5['Desvio Tempo em %'] < -30) | (azp5['Desvio Tempo em %'] > 30)]

                #Repetitvo
                excluir_2 = [3,4,5]
                repetitivo_1 = azp2[(azp2['Desvio Tempo em %'] < -30) | (azp2['Desvio Tempo em %'] > 30)]
                df = pd.DataFrame(repetitivo_1)
                df = df.drop(df.columns[excluir_2], axis=1)
                df_repetitivo_1 = df.dropna(how='all', axis=0, inplace=False)
                
                #Diario
                diario_1 = azp3[(azp3['Desvio Tempo em %'] < -30) | (azp3['Desvio Tempo em %'] > 30)]
                df = pd.DataFrame(diario_1)

                df_diario_1 = df.dropna(how='all', axis=0, inplace=False)


                #semanal
                df = pd.DataFrame(semana_1)
                df_semana_1 = df.dropna(how='all', axis=0, inplace=False)


                #Mensal
                df = pd.DataFrame(mensal_1)
                df_mensal_1 = df.dropna(how='all', axis=0, inplace=False)

                # 1.7.	Jobs com atraso no início da execução (>10min)         

                atraso = bd_azp1.loc[bd_azp1['Atraso ( segs )'] > 60000 ]

                df = pd.DataFrame(atraso)
                df_atraso = df.dropna(how='all', axis=0, inplace=False)

                # Adicione um título ao documento
                doc.add_heading('Sumário', level=1)

                # Crie o sumário
                sumario_para_adicionar = '\n'.join(sumario)
                doc.add_paragraph(sumario_para_adicionar)

                doc.add_paragraph('')
                doc.add_heading('1.AZP', level=0)
                
                
                # Jobs Cancelados
                paragrafo1 = doc.add_paragraph('1.1 Jobs Cancelados\n')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph(f'Cancelamento de Jobs no periodo de {periodo_inicio} às 01:00hrs {periodo_final} até 00:59hrs.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_cancelado, 'Cancelados(1.1 AZP)')
                doc.add_paragraph('')
                # 1.2 Jobs Previstos e não executados
                paragrafo1 = doc.add_paragraph('1.2 Jobs Previstos e não executados\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_repetitivo, 'Repetitivo(1.2 AZP)')

                add_dataframe_as_table(db_diario, 'Diário(1.2 AZP)')

                add_dataframe_as_table(df_semana, 'Semanal(1.2 AZP)')

                add_dataframe_as_table(df_mensal, 'Mensal(1.2 AZP)')
                doc.add_paragraph('')

                #1.3 Desvio de ocorrências previstas x executadas
                paragrafo1 = doc.add_paragraph('1.3 Desvio de ocorrências previstas x executadas\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(desvio_ocorrencias, 'Repetitivo(1.3 AZP)')
                doc.add_paragraph('')

                #1.4 Jobs Ativos Duplicados
                paragrafo1 = doc.add_paragraph('1.4 Jobs Ativos Duplicados')
                add_dataframe_as_table(azp6, '1.4 Jobs Ativos(1.4 AZP)') 
                doc.add_paragraph('')
                #1.5 Jobs Ativos Acima de 24h
                paragrafo1 = doc.add_paragraph('1.5 Jobs Ativos Acima de 24h')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])

                #1.6.	Jobs com tempo de execução diferente do previsto
                paragrafo1 = doc.add_paragraph('1.6.Jobs com tempo de execução diferente do previsto\n')
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph("Identificados Jobs com desvio de execução com impacto significativo (acima de 30%)\n")

                add_dataframe_as_table(df_repetitivo, 'Repetitivo(1.6 AZP)')

                add_dataframe_as_table(df_diario_1, 'Diário(1.6 AZP)')

                add_dataframe_as_table(df_semana_1, 'Semanal(1.6 AZP)')

                add_dataframe_as_table(df_mensal_1, 'Mensal(1.6 AZP)')
                doc.add_paragraph('')

                # 1.7.	Jobs com atraso no início da execução (>10min)

                paragrafo1 = doc.add_paragraph('1.7.	Jobs com atraso no início da execução (>10min)\n')
                doc.add_paragraph(f'Jobs com atraso acima de 10 minutos.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])

                add_dataframe_as_table(df_atraso, 'Jobs com Atraso(1.7 AZP)')
                doc.add_paragraph('')

                paragrafo1 = doc.add_paragraph('1.8 DUMPS - AZP')
                aplicar_formatacao(paragrafo1.runs[0])
                
                add_dataframe_as_table(bd_azp7, 'DUMPS (1.8 AZP)')
                doc.add_paragraph('')

                #-----------------------AMBIENTE FIP----------------------#
                doc.add_heading('2.FIP', level=0)
                doc.add_paragraph('')
                
                # Tabelas e filtros por base

                bd_fip1 = pd.read_excel(banco_fip ,sheet_name = 'Jobs')
                bd_fip2 = pd.read_excel(banco_fip,sheet_name = 'Repetitivo')
                bd_fip3 = pd.read_excel(banco_fip,sheet_name = 'Diário')
                bd_fip4 = pd.read_excel(banco_fip,sheet_name = 'Semanal')
                bd_fip5 = pd.read_excel(banco_fip,sheet_name = 'Mensal')
                bd_fip6 = pd.read_excel(banco_fip,sheet_name = 'Ativo_24h')
                bd_fip7 = pd.read_excel(banco_fip,sheet_name = 'DUMPS')
                
                fip2 = bd_fip2.iloc[ : , 0:11]
                fip3 = bd_fip3.iloc[ : , 0:9]
                fip4 = bd_fip4.iloc[ : , 0:11]
                fip5 = bd_fip5.iloc[ : , 0:11]
                fip6 = bd_fip6.iloc[0:8]

                cancelado = bd_fip1.loc[bd_fip1['Status'] == 'cancelado']
                df = pd.DataFrame(cancelado)
                df_fip_cancelado = df.dropna(how='all', axis=0, inplace=False)
                

                # 2.2 Jobs Previstos e não executados
                #Repetitvo        
                repetitivo = fip2.iloc[:, :11][fip2['Ocorrencias Reais'] == 0]
                df = pd.DataFrame(repetitivo)
                # Converter a coluna diretamente para inteiros (se não houver valores inválidos ou nulos)
                df_fip_repetitivo = df.dropna(how='all', axis=0, inplace=False)
                

                #Diario
                diario = fip3.iloc[:, :10][fip3['Status'] == 'cancelado']
                df = pd.DataFrame(diario)
                df_fip_diario = df.dropna(how='all', axis=0, inplace=False)
                

                #semanal
                semana = fip4.iloc[:, :10][fip4['Dia da Semana'] == dia_da_semana_de_ontem_formatado]
                df = pd.DataFrame(semana)
                df_fip_semana = df.dropna(how='all', axis=0, inplace=False)

                #Mensal
                mensal = fip5.iloc[:, :9][fip5['Dia do Mes'] == dia_mes]
                df = pd.DataFrame(mensal)
                df_fip_mensal = df.dropna(how='all', axis=0, inplace=False)
                

                #2.3 Desvio de ocorrências previstas x executadas
                                
                #repetitivo_0 = tabela2.loc[tabela2['Desvio Tempo em % '] != 0, ['ID', 'Tipo','Nomenclatura / Nome do Job','Frequencia (Minutos)', 'Ocorrencias Previstas', 'Ocorrencias Reais','Desvio Tempo em % ','Status Atual']]
                
                repetitivo_0 = fip2.iloc[:, :11][fip2['Desvio Tempo em %'] != 0]
                repetitivo_0 = repetitivo_0.dropna(subset=['Desvio Tempo em %'])

                df = pd.DataFrame(repetitivo_0)
                
                # Convertendo a coluna 'Desvio Tempo em %' para tipo numérico
                df['Desvio Tempo em %'] = pd.to_numeric(df['Desvio Tempo em %'], errors='coerce')
                # Arredonda os valores na coluna 'Desvio Tempo em %' para duas casas decimais
                df['Desvio Tempo em %'] = df['Desvio Tempo em %'].round(2)
                # Convertendo a coluna 'Desvio Tempo em %' para tipo numérico
                df['Tempo Real Execução'] = pd.to_numeric(df['Tempo Real Execução'], errors='coerce')
                # Arredonda os valores na coluna 'Desvio Tempo em %' para duas casas decimais
                df['Tempo Real Execução'] = df['Tempo Real Execução'].round(2)


                df_fip_desvio_ocorrecia = df.dropna(how='all', axis=0, inplace=False)
                #2.4 Jobs Ativos Duplicados
                aplicar_formatacao(paragrafo1.runs[0])
                
                
                #2.5 Jobs Ativos Acima de 24h
                            
                #2.6.	Jobs com tempo de execução diferente do previsto
                repetitivo_1 = fip2[(fip2['Desvio Tempo em %'] < -30) | (fip2['Desvio Tempo em %'] > 30)]
                diario_1 = fip3[(fip3['Desvio Tempo em %'] < -30) | (fip3['Desvio Tempo em %'] > 30)]
                semana_1 = fip4[(fip4['Desvio Tempo em %'] < -30) | (fip4['Desvio Tempo em %'] > 30)]
                mensal_1 = fip5[(fip5['Desvio Tempo em %'] < -30) | (fip5['Desvio Tempo em %'] > 30)]

                #Repetitvo
                            
                df = pd.DataFrame(repetitivo_1)
                # Convertendo a coluna 'Desvio Tempo em %' para tipo numérico
                df['Desvio Tempo em %'] = pd.to_numeric(df['Desvio Tempo em %'], errors='coerce')
                # Arredonda os valores na coluna 'Desvio Tempo em %' para duas casas decimais
                df['Desvio Tempo em %'] = df['Desvio Tempo em %'].round(2)
                # Convertendo a coluna 'Desvio Tempo em %' para tipo numérico
                df['Tempo Real Execução'] = pd.to_numeric(df['Tempo Real Execução'], errors='coerce')
                # Arredonda os valores na coluna 'Desvio Tempo em %' para duas casas decimais
                df['Tempo Real Execução'] = df['Tempo Real Execução'].round(2)
                df_fip_repetitivo1 = df.dropna(how='all', axis=0, inplace=False)
                

                #Diario
                df = pd.DataFrame(diario_1)
                # Convertendo a coluna 'Desvio Tempo em %' para tipo numérico
                df['Desvio Tempo em %'] = pd.to_numeric(df['Desvio Tempo em %'], errors='coerce')
                # Arredonda os valores na coluna 'Desvio Tempo em %' para duas casas decimais
                df['Desvio Tempo em %'] = df['Desvio Tempo em %'].round(2)
                df_fip_diario1 = df.dropna(how='all', axis=0, inplace=False)
                

                #semanal
                df = pd.DataFrame(semana_1)
                df_fip_semana1 = df.dropna(how='all', axis=0, inplace=False)
                

                #Mensal
                df = pd.DataFrame(mensal_1)
                df_fip_mensal1 = df.dropna(how='all', axis=0, inplace=False)
                

                # 2.7.	Jobs com atraso no início da execução (>10min)         
                
                atraso = bd_fip1.loc[bd_fip1['Atraso ( segs )'] >= 60000 ]
                df = pd.DataFrame(atraso)
                df_fip_atraso = df.dropna(how='all', axis=0, inplace=False)
            

                # Jobs Cancelados
                paragrafo1 = doc.add_paragraph('2.1 Jobs Cancelados\n')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph(f'Cancelamento de Jobs no periodo de {periodo_inicio} às 01:00hrs {periodo_final} até 00:59hrs.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_fip_cancelado, 'Cancelados(2.1 FIP)')
                doc.add_paragraph('')
                # 1.2 Jobs Previstos e não executados
                paragrafo1 = doc.add_paragraph('2.2 Jobs Previstos e não executados\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_fip_repetitivo, 'Repetitivo(2.2 FIP)')

                add_dataframe_as_table(df_fip_diario, 'Diário(2.2 FIP)')

                add_dataframe_as_table(df_fip_semana, 'Semanal(2.2 FIP)')

                add_dataframe_as_table(df_fip_mensal, 'Mensal(2.2 FIP)')
                doc.add_paragraph('')

                #2.3 Desvio de ocorrências previstas x executadas
                paragrafo1 = doc.add_paragraph('2.3 Desvio de ocorrências previstas x executadas\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_fip_desvio_ocorrecia, 'Repetitivo(2.3 FIP)')
                doc.add_paragraph('')

                #2.4 Jobs Ativos Duplicados
                paragrafo1 = doc.add_paragraph('2.4 Jobs Ativos Duplicados')
                add_dataframe_as_table(fip6, '2.4 Jobs Ativos(2.4 FIP)')   
                doc.add_paragraph('') 
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])

                #2.5 Jobs Ativos Acima de 24h
                paragrafo1 = doc.add_paragraph('2.5 Jobs Ativos Acima de 24h')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])

                #2.6.	Jobs com tempo de execução diferente do previsto
                paragrafo1 = doc.add_paragraph('2.6.Jobs com tempo de execução diferente do previsto\n')
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph("Identificados Jobs com desvio de execução com impacto significativo (acima de 30%)\n")

                add_dataframe_as_table(df_fip_repetitivo1, 'Repetitivo(2.6 FIP)')

                add_dataframe_as_table(df_fip_diario1, 'Diário(2.6 FIP)')

                add_dataframe_as_table(df_fip_semana1, 'Semanal(2.6 FIP)')

                add_dataframe_as_table(df_fip_mensal1, 'Mensal(2.6 FIP)')
                doc.add_paragraph('')

                # 2.7.	Jobs com atraso no início da execução (>10min)

                paragrafo1 = doc.add_paragraph('2.7.	Jobs com atraso no início da execução (>10min)\n')
                doc.add_paragraph(f'Jobs com atraso acima de 10 minutos' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])

                add_dataframe_as_table(df_fip_atraso, 'Jobs com Atraso(2.7 FIP)')
                doc.add_paragraph('')

                #---------DUMPS-------------#
                paragrafo1 = doc.add_paragraph('2.8 DUMPS - FIP')
                aplicar_formatacao(paragrafo1.runs[0])
                
                add_dataframe_as_table(bd_fip7, 'DUMPS (2.8 FIP)')
                doc.add_paragraph('')

                #---------------PIP----------------#
                
                doc.add_heading('3.PIP', level=0)
                doc.add_paragraph('')
                        
                # Tabelas e filtros por base

                bd_pip1 = pd.read_excel(banco_pip,sheet_name = 'Jobs')
                bd_pip2 = pd.read_excel(banco_pip,sheet_name = 'Repetitivo')
                bd_pip3 = pd.read_excel(banco_pip,sheet_name = 'Diário')
                bd_pip4 = pd.read_excel(banco_pip,sheet_name = 'Semanal')
                bd_pip5 = pd.read_excel(banco_pip,sheet_name = 'Mensal')
                bd_pip6 = pd.read_excel(banco_pip,sheet_name = 'Ativo_24h')
                bd_pip7 = pd.read_excel(banco_pip,sheet_name = 'DUMPS')

                pip2 = bd_pip2.iloc[ : , 0:11]
                pip3 = bd_pip3.iloc[ : , 0:9]
                pip4 = bd_pip4.iloc[ : , 0:11]
                pip5 = bd_pip5.iloc[ : , 0:11]
                pip6 = bd_pip6.iloc[0:8]
            

                cancelado = bd_pip1.loc[bd_pip1['Status'] == 'cancelado']
                df = pd.DataFrame(cancelado)
                df_pip_cancelado = df.dropna(how='all', axis=0, inplace=False)
                

                # 3.2 Jobs Previstos e não executados
                #Repetitvo        
                repetitivo = pip2.iloc[:, :11][pip2['Ocorrencias Reais'] == 0]
                df = pd.DataFrame(repetitivo)
                df_pip_repetitivo = df.dropna(how='all', axis=0, inplace=False)
                

                #Diario
                diario = pip3.iloc[:,0:10][pip3['Status'] == 'cancelado']
                df = pd.DataFrame(diario)
                df_pip_diario = df.dropna(how='all', axis=0, inplace=False)
                

                #semanal
                semana = pip4.iloc[:, 0:11][pip4['Dia da Semana'] == dia_da_semana_de_ontem_formatado]
                df = pd.DataFrame(semana)
                df_pip_semana = df.dropna(how='all', axis=0, inplace=False)

                #Mensal
                mensal = pip5.iloc[:, :11][pip5['Dia do Mes'] == dia_mes]
                df = pd.DataFrame(mensal)
                df_pip_mensal = df.dropna(how='all', axis=0, inplace=False)
                

                #3.3 Desvio de ocorrências previstas x executadas
                                
                #repetitivo_0 = tabela2.loc[tabela2['Desvio Tempo em % '] != 0, ['ID', 'Tipo','Nomenclatura / Nome do Job','Frequencia (Minutos)', 'Ocorrenci23as Previstas', 'Ocorrencias Reais','Desvio Tempo em % ','Status Atual']]
                repetitivo_0 = pip2.iloc[:, 0:11][pip2['Desvio Tempo em %'] != 0]
                repetitivo_0 = repetitivo_0.dropna(subset=['Desvio Tempo em %'])
                df = pd.DataFrame(repetitivo_0)
                df_pip_desvio_ocorrecia = df.dropna(how='all', axis=0, inplace=False)
                

                #3.4 Jobs Ativos Duplicados
                df = pd.DataFrame(pip6)
                df_pip_semana = df.dropna(how='all', axis=0, inplace=False)   

                #3.5 Jobs Ativos Acima de 24h
                            
                #3.6.	Jobs com tempo de execução diferente do previsto
                repetitivo_1 = pip2[(pip2['Desvio Tempo em %'] < -30) | (pip2['Desvio Tempo em %'] > 30)]
                diario_1 = pip3[(pip3['Desvio Tempo em %'] < -30) | (pip3['Desvio Tempo em %'] > 30)]
                semana_1 = pip4[(pip4['Desvio Tempo em %'] < -30) | (pip4['Desvio Tempo em %'] > 30)]
                mensal_1 = pip5[(pip5['Desvio Tempo em %'] < -30) | (pip5['Desvio Tempo em %'] > 30)]

                #Repetitvo
                            
                df = pd.DataFrame(repetitivo_1)
                df_pip_repetitivo1 = df.dropna(how='all', axis=0, inplace=False)
                

                #Diario
                df = pd.DataFrame(diario_1)
                # Convertendo a coluna 'Desvio Tempo em %' para tipo numérico
                df['Desvio Tempo em %'] = pd.to_numeric(df['Desvio Tempo em %'], errors='coerce')
                # Arredonda os valores na coluna 'Desvio Tempo em %' para duas casas decimais
                df['Desvio Tempo em %'] = df['Desvio Tempo em %'].round(2)
                df_pip_diario1 = df.dropna(how='all', axis=0, inplace=False)
                

                #semanal
                df = pd.DataFrame(semana_1)
                df_pip_semana1 = df.dropna(how='all', axis=0, inplace=False)
                

                #Mensal
                df = pd.DataFrame(mensal_1)
                df_pip_mensal1 = df.dropna(how='all', axis=0, inplace=False)
                

                # 3.7.	Jobs com atraso no início da execução (>10min)         
                
                atraso = bd_pip1.loc[bd_pip1['Atraso ( segs )'] >= 60000 ]
                df = pd.DataFrame(atraso)
                df_pip_atraso = df.dropna(how='all', axis=0, inplace=False)
            

                # Jobs Cancelados
                paragrafo1 = doc.add_paragraph('3.1 Jobs Cancelados\n')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph(f'Cancelamento de Jobs no periodo de {periodo_inicio} às 01:00hrs {periodo_final} até 00:59hrs.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_pip_cancelado, 'Cancelados(3.1 PIP)')
                doc.add_paragraph('')
                # 3.2 Jobs Previstos e não executados
                paragrafo1 = doc.add_paragraph('3.2 Jobs Previstos e não executados\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_pip_repetitivo, 'Repetitivo(3.2 PIP)')

                add_dataframe_as_table(df_pip_diario, 'Diário(3.2 PIP)')

                add_dataframe_as_table(semana, 'Semanal(3.2 PIP)')

                add_dataframe_as_table(df_pip_mensal, 'Mensal(3.2 PIP)')
                doc.add_paragraph('')

                #3.3 Desvio de ocorrências previstas x executadas
                paragrafo1 = doc.add_paragraph('3.3 Desvio de ocorrências previstas x executadas\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_pip_desvio_ocorrecia, 'Repetitivo(3.3 PIP)')
                doc.add_paragraph('')

                #3.4 Jobs Ativos Duplicados
                paragrafo1 = doc.add_paragraph('3.4 Jobs Ativos Duplicados')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(pip6, '3.4 Jobs Ativos(3.4 PIP)')
                doc.add_paragraph('')
                #3.5 Jobs Ativos Acima de 24h
                paragrafo1 = doc.add_paragraph('3.5 Jobs Ativos Acima de 24h')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])

                #3.6.	Jobs com tempo de execução diferente do previsto
                paragrafo1 = doc.add_paragraph('3.6.Jobs com tempo de execução diferente do previsto\n')
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph("Identificados Jobs com desvio de execução com impacto significativo (acima de 30%)\n")

                add_dataframe_as_table(df_pip_repetitivo1, 'Repetitivo(3.6 PIP)')

                add_dataframe_as_table(df_pip_diario1, 'Diário(3.6 PIP)')

                add_dataframe_as_table(df_pip_semana1, 'Semanal(3.6 PIP)')

                add_dataframe_as_table(df_pip_mensal1, 'Mensal(3.6 PIP)')
                doc.add_paragraph('')

                #3.7.	Jobs com atraso no início da execução (>10min)

                paragrafo1 = doc.add_paragraph('3.7.	Jobs com atraso no início da execução (>10min)\n')
                doc.add_paragraph(f'Jobs com atraso acima de 10 minutos.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])

                add_dataframe_as_table(df_pip_atraso, 'Jobs com Atraso((3.7 PIP))')
                doc.add_paragraph('')


                #-------DUMPS----------#
                paragrafo1 = doc.add_paragraph('3.8 DUMPS - PIP')
                aplicar_formatacao(paragrafo1.runs[0])
                
                add_dataframe_as_table(bd_pip7, 'DUMPS (3.8 PIP)')
                doc.add_paragraph('')


                #-----------------GRP---------------------#

                doc.add_heading('4.GRP', level=0)
                
                doc.add_paragraph('')
                #Tabelas e filtros por base
            
                bd_grp1 = pd.read_excel(banco_grp ,sheet_name = 'Jobs')
                bd_grp2 = pd.read_excel(banco_grp,sheet_name = 'Repetitivo')
                bd_grp3 = pd.read_excel(banco_grp,sheet_name = 'Diário')
                bd_grp4 = pd.read_excel(banco_grp,sheet_name = 'Semanal')
                bd_grp5 = pd.read_excel(banco_grp,sheet_name = 'Mensal')
                bd_grp6 = pd.read_excel(banco_grp ,sheet_name = 'Ativo_24h')
                bd_grp7 = pd.read_excel(banco_grp,sheet_name = 'DUMPS')

                grp2 = bd_grp2.iloc[ : , 0:11]
                grp3 = bd_grp3.iloc[ : , 0:9]
                grp4 = bd_grp4.iloc[ : , 0:11]
                grp5 = bd_grp5.iloc[ : , 0:11]
                grp6 = bd_grp6.iloc[0:8]

                cancelado = bd_grp1.loc[bd_grp1['Status'] == 'cancelado']
                df = pd.DataFrame(cancelado)
                df_grp_cancelado = df.dropna(how='all', axis=0, inplace=False)
                

                # 4.2 Jobs Previstos e não executados
                #Repetitvo        
                repetitivo = grp2.iloc[:, :11][grp2['Ocorrencias Reais'] == 0]
                df = pd.DataFrame(repetitivo)
                df_grp_repetitivo = df.dropna(how='all', axis=0, inplace=False)
                

                #Diario
                diario = grp3.iloc[:, :10][grp3['Status'] == 'cancelado']
                df = pd.DataFrame(diario)
                df_grp_diario = df.dropna(how='all', axis=0, inplace=False)
                

                #semanal
                semana = grp4.iloc[:, :10][grp4['Dia da Semana'] == dia_da_semana_de_ontem_formatado]
                df = pd.DataFrame(semana)
                df_grp_semana = df.dropna(how='all', axis=0, inplace=False)

                #Mensal
                mensal = grp5.iloc[:, :9][grp5['Dia do Mes'] == dia_mes]
                df = pd.DataFrame(mensal)
                df_grp_mensal = df.dropna(how='all', axis=0, inplace=False)
                

                #4.3 Desvio de ocorrências previstas x executadas
                                
                #repetitivo_0 = tabela2.loc[tabela2['Desvio Tempo em % '] != 0, ['ID', 'Tipo','Nomenclatura / Nome do Job','Frequencia (Minutos)', 'Ocorrencias Previstas', 'Ocorrencias Reais','Desvio Tempo em % ','Status Atual']]
                repetitivo_0 = grp2.iloc[:, :11][grp2['Desvio Tempo em %'] != 0]
                df_grp_desvio_ocorrecia = repetitivo_0.dropna(subset=['Desvio Tempo em %'])            
                df = pd.DataFrame(repetitivo_0)
                
                

                #4.4 Jobs Ativos Duplicados
                df = pd.DataFrame(pip6)
                df_pip_semana = df.dropna(how='all', axis=0, inplace=False)
                #4.5 Jobs Ativos Acima de 24h
                            
                #4.6.	Jobs com tempo de execução diferente do previsto
                repetitivo_1 = grp2[(grp2['Desvio Tempo em %'] < -30) | (grp2['Desvio Tempo em %'] > 30)]
                diario_1 = grp3[(grp3['Desvio Tempo em %'] < -30) | (grp3['Desvio Tempo em %'] > 30)]
                semana_1 = grp4[(grp4['Desvio Tempo em %'] < -30) | (grp4['Desvio Tempo em %'] > 30)]
                mensal_1 = grp5[(grp5['Desvio Tempo em %'] < -30) | (grp5['Desvio Tempo em %'] > 30)]

                #Repetitvo
                            
                df = pd.DataFrame(repetitivo_1)
                df_grp_repetitivo1 = df.dropna(how='all', axis=0, inplace=False)
                

                #Diario
                df = pd.DataFrame(diario_1)
                df_grp_diario1 = df.dropna(how='all', axis=0, inplace=False)
                

                #semanal
                df = pd.DataFrame(semana_1)
                df_grp_semana1 = df.dropna(how='all', axis=0, inplace=False)
                

                #Mensal
                df = pd.DataFrame(mensal_1)
                df_grp_mensal1 = df.dropna(how='all', axis=0, inplace=False)
                

                # 4.7.	Jobs com atraso no início da execução (>10min)         
                
                atraso = bd_grp1.loc[bd_grp1['Atraso ( segs )'] >= 60000 ]
                df = pd.DataFrame(atraso)
                df_grp_atraso = df.dropna(how='all', axis=0, inplace=False)
            

                # Jobs Cancelados
                paragrafo1 = doc.add_paragraph('4.1 Jobs Cancelados\n')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph(f'Cancelamento de Jobs no periodo de {periodo_inicio} às 01:00hrs {periodo_final} até 00:59hrs.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_grp_cancelado, 'Cancelados(4.1 GRP)')
                doc.add_paragraph('')
                #4.2 Jobs Previstos e não executados
                paragrafo1 = doc.add_paragraph('4.2 Jobs Previstos e não executados\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_grp_repetitivo, 'Repetitivo(4.2 GRP)')

                add_dataframe_as_table(df_grp_diario, 'Diário(4.2 GRP)')

                add_dataframe_as_table(df_grp_semana, 'Semanal(4.2 GRP)')

                add_dataframe_as_table(df_grp_mensal, 'Mensal(4.2 GRP)')
                doc.add_paragraph('')

                #4.3 Desvio de ocorrências previstas x executadas
                paragrafo1 = doc.add_paragraph('4.3 Desvio de ocorrências previstas x executadas\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_grp_desvio_ocorrecia, 'Repetitivo(4.3 GRP)')
                doc.add_paragraph('')

                #4.4 Jobs Ativos Duplicados
                paragrafo1 = doc.add_paragraph('4.4 Jobs Ativos Duplicados')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(grp6, '4.4 Jobs Ativos(4.4 PIP)')
                #4.5 Jobs Ativos Acima de 24h
                paragrafo1 = doc.add_paragraph('4.5 Jobs Ativos Acima de 24h')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph('')

                #4.6.	Jobs com tempo de execução diferente do previsto
                paragrafo1 = doc.add_paragraph('4.6.Jobs com tempo de execução diferente do previsto\n')
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph("Identificados Jobs com desvio de execução com impacto significativo (acima de 30%)\n")

                add_dataframe_as_table(df_grp_repetitivo1, 'Repetitivo(4.6 GRP)')

                add_dataframe_as_table(df_grp_diario1, 'Diário(4.6 GRP)')

                add_dataframe_as_table(df_grp_semana1, 'Semanal(4.6 GRP)')

                add_dataframe_as_table(df_grp_mensal1, 'Mensal(4.6 GRP)')
                doc.add_paragraph('')

                #4.7.	Jobs com atraso no início da execução (>10min)

                paragrafo1 = doc.add_paragraph('4.7.	Jobs com atraso no início da execução (>10min)\n')
                doc.add_paragraph(f'Jobs com atraso acima de 10 minutos.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_grp_atraso, 'Jobs com Atraso(4.7 GRP)')
                doc.add_paragraph('') 

                #------DUMPS----------#
                paragrafo1 = doc.add_paragraph('4.8 DUMPS - GRP')
                aplicar_formatacao(paragrafo1.runs[0])
                
                add_dataframe_as_table(bd_grp7, 'DUMPS (4.8 GRP)')
                doc.add_paragraph('')
                
                #----------BIP-------------#

                doc.add_heading('5.BIP', level=0)
                doc.add_paragraph('')
                #Tabelas e filtros por base
                
                bd_bip1 = pd.read_excel(banco_bip ,sheet_name = 'Jobs')
                bd_bip2 = pd.read_excel(banco_bip,sheet_name = 'Repetitivo')
                bd_bip3 = pd.read_excel(banco_bip,sheet_name = 'Diário')
                bd_bip4 = pd.read_excel(banco_bip,sheet_name = 'Semanal')
                bd_bip5 = pd.read_excel(banco_bip,sheet_name = 'Mensal')
                bd_bip6 = pd.read_excel(banco_bip,sheet_name = 'Ativo_24h')
                bd_bip7 = pd.read_excel(banco_bip,sheet_name = 'DUMPS')    

                bip2 = bd_bip2.iloc[ : , 0:11]
                bip3 = bd_bip3.iloc[ : , 0:9]
                bip4 = bd_bip4.iloc[ : , 0:11]
                bip5 = bd_bip5.iloc[ : , 0:11]
                bip6 = bd_bip6.iloc[0:8]

                cancelado = bd_bip1.loc[bd_bip1['Status'] == 'cancelado']
                df = pd.DataFrame(cancelado)
                df_bip_cancelado = df.dropna(how='all', axis=0, inplace=False)
                

                # 5.2 Jobs Previstos e não executados
                #Repetitvo        
                repetitivo = bip2.iloc[:, :10][bip2['Ocorrencias Reais'] == 0]
                df = pd.DataFrame(repetitivo)
                df_bip_repetitivo = df.dropna(how='all', axis=0, inplace=False)
                

                #Diario
                diario = bip3.iloc[:, :10][bip3['Status'] == 'cancelado']
                df = pd.DataFrame(diario)
                df_bip_diario = df.dropna(how='all', axis=0, inplace=False)
                

                #semanal
                semana = bip4.iloc[:, :10][bip4['Dia da Semana'] == dia_da_semana_de_ontem_formatado]
                df = pd.DataFrame(semana)
                df_bip_semana = df.dropna(how='all', axis=0, inplace=False)

                #Mensal
                mensal = bip5.iloc[:, :9][bip5['Dia do Mes'] == dia_mes]
                df = pd.DataFrame(mensal)
                df_bip_mensal = df.dropna(how='all', axis=0, inplace=False)
                
                #5.3 Desvio de ocorrências previstas x executadas
                                
                #repetitivo_0 = tabela2.loc[tabela2['Desvio Tempo em % '] != 0, ['ID', 'Tipo','Nomenclatura / Nome do Job','Frequencia (Minutos)', 'Ocorrencias Previstas', 'Ocorrencias Reais','Desvio Tempo em % ','Status Atual']]
                repetitivo_0 = bip2.iloc[:, :11][bip2['Desvio Tempo em %'] != 0]
                df = pd.DataFrame(repetitivo_0)
                df_bip_desvio_ocorrecia = df.dropna(subset=['Desvio Tempo em %']).round(2)
                
                
                #5.4 Jobs Ativos Duplicados
                df = pd.DataFrame(bip6)
                #5.5 Jobs Ativos Acima de 24h
                            
                #5.6.	Jobs com tempo de execução diferente do previsto
                repetitivo_1 = bip2[(bip2['Desvio Tempo em %'] < -30) | (bip2['Desvio Tempo em %'] > 30)]
                diario_1 = bip3[(bip3['Desvio Tempo em %'] < -30) | (bip3['Desvio Tempo em %'] > 30)]
                semana_1 = bip4[(bip4['Desvio Tempo em %'] < -30) | (bip4['Desvio Tempo em %'] > 30)]
                mensal_1 = bip5[(bip5['Desvio Tempo em %'] < -30) | (bip5['Desvio Tempo em %'] > 30)]

                #Repetitvo
                    
                df = pd.DataFrame(repetitivo_1)
                df_bip_repetitivo1 = df.dropna(how='all', axis=0, inplace=False).round(2)
                #df_bip_repetitivo1.round(2)

                #Diario
                df = pd.DataFrame(diario_1)
                df_bip_diario1 = df.dropna(how='all', axis=0, inplace=False).round(2)
                

                #semanal
                df = pd.DataFrame(semana_1)
                df_bip_semana1 = df.dropna(how='all', axis=0, inplace=False).round(2)
                

                #Mensal
                df = pd.DataFrame(mensal_1)
                df_bip_mensal1 = df.dropna(how='all', axis=0, inplace=False).round(2)
                        

                # 5.7.	Jobs com atraso no início da execução (>10min)         
                
                atraso = bd_bip1.loc[bd_bip1['Atraso ( segs )'] >= 60000 ]
                df = pd.DataFrame(atraso)
                df_bip_atraso = df.dropna(how='all', axis=0, inplace=False).round(2)
            

                # Jobs Cancelados
                paragrafo1 = doc.add_paragraph('5.1 Jobs Cancelados\n')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph(f'Cancelamento de Jobs no periodo de {periodo_inicio} às 01:00hrs {periodo_final} até 00:59hrs.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_bip_cancelado, 'Cancelados(5.1 BIP)')
                doc.add_paragraph('')
                # 5.2 Jobs Previstos e não executados
                paragrafo1 = doc.add_paragraph('5.2 Jobs Previstos e não executados\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_bip_repetitivo, 'Repetitivo(5.2 BIP)')

                add_dataframe_as_table(df_bip_diario, 'Diário(5.2 BIP)')

                add_dataframe_as_table(df_bip_semana, 'Semanal(5.2 BIP)')

                add_dataframe_as_table(df_bip_mensal, 'Mensal(5.2 BIP)')
                doc.add_paragraph('')

                #5.3 Desvio de ocorrências previstas x executadas
                paragrafo1 = doc.add_paragraph('5.3 Desvio de ocorrências previstas x executadas\n')
                doc.add_paragraph(f'Todos os Jobs previstos que não foram executados.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(df_bip_desvio_ocorrecia, 'Repetitivo(5.3 BIP)')
                doc.add_paragraph('')

                #5.4 Jobs Ativos Duplicados
                paragrafo1 = doc.add_paragraph('5.4 Jobs Ativos Duplicados')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])
                add_dataframe_as_table(bip6, '5.4 Jobs Ativos(5.4 BIP)')
            
                #5.5 Jobs Ativos Acima de 24h
                paragrafo1 = doc.add_paragraph('5.5 Jobs Ativos Acima de 24h')
                doc.add_paragraph('')
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])

                #5.6.	Jobs com tempo de execução diferente do previsto
                paragrafo1 = doc.add_paragraph('5.6.Jobs com tempo de execução diferente do previsto\n')
                aplicar_formatacao(paragrafo1.runs[0])
                doc.add_paragraph("Identificados Jobs com desvio de execução com impacto significativo (acima de 30%)\n")

                add_dataframe_as_table(df_bip_repetitivo1, 'Repetitivo(5.6 BIP)')

                add_dataframe_as_table(df_bip_diario1, 'Diário(5.6 BIP)')

                add_dataframe_as_table(df_bip_semana1, 'Semanal(5.6 BIP)')

                add_dataframe_as_table(df_bip_mensal1, 'Mensal(5.6 BIP)')
                doc.add_paragraph('')

                # 5.7.	Jobs com atraso no início da execução (>10min)

                paragrafo1 = doc.add_paragraph('5.7.	Jobs com atraso no início da execução (>10min)\n')
                doc.add_paragraph(f'Jobs com atraso acima de 10 minutos.' )
                # Acesse a execução (run) dentro do parágrafo e aplique o padrão de formatação
                aplicar_formatacao(paragrafo1.runs[0])

                add_dataframe_as_table(df_bip_atraso, 'Jobs com Atraso(5.7 BIP)')
                doc.add_paragraph('') 
                
                #-------DUMPS--------------#  
                paragrafo1 = doc.add_paragraph('5.8 DUMPS - BIP')
                aplicar_formatacao(paragrafo1.runs[0])
                
                add_dataframe_as_table(bd_bip7, 'DUMPS (5.8 BIP)')
                doc.add_paragraph('')
                            
                st.write("Processo finalizado com sucesso!")
                
                # Salva o documento em um buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                
                # Move o buffer para o início
                buffer.seek(0)
                
                # Botão de download
                st.download_button(
                    label="Baixar Documento",
                    data=buffer,
                    file_name=f"Monitoramento de Dumps e Jobs dos ambientes SAP {periodo_inicio} a {periodo_final}.doc",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # Rodapé
    st.markdown('---')
    st.markdown('© 2024 Desenvolvido por Robson Pereira.')

        # Incluir o CSS personalizado usando o método st.markdown
if __name__ == "__main__":
    sap()   
